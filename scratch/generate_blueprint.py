import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_blueprint():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    styles = doc.styles
    
    # Colors
    NAVY = RGBColor(12, 28, 48)      # #0C1C30 - Primary
    GOLD = RGBColor(197, 160, 89)     # #C5A059 - Tertiary / Gold Accent
    CHARCOAL = RGBColor(51, 51, 51)   # #333333 - Body Text
    MUTED_GRAY = RGBColor(115, 115, 115) # Muted Gray
    
    # Document Title (Cover Page Style)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("REGENT ADVISORY\n")
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = GOLD
    
    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_p2.add_run("PROMPT-DRIVEN AI ADVISORY CHATBOT\nMASTER BLUEPRINT")
    title_run2.font.name = 'Georgia'
    title_run2.font.size = Pt(26)
    title_run2.font.bold = True
    title_run2.font.color.rgb = NAVY
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Strategic Design & Implementation Plan for Next-Generation Luxury Real Estate Client Engagement")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = CHARCOAL
    
    doc.add_paragraph("\n" * 2)
    
    # Metadata Box
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    cell = meta_table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=200, bottom=200, left=300, right=300)
    
    meta_p = cell.paragraphs[0]
    meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run1 = meta_p.add_run("Prepared For: ")
    meta_run1.font.bold = True
    meta_run1.font.name = 'Arial'
    meta_run1.font.size = Pt(10)
    meta_run1.font.color.rgb = NAVY
    
    meta_p.add_run("Regent Advisory Firm Partners\n")
    
    meta_run2 = meta_p.add_run("Version: ")
    meta_run2.font.bold = True
    meta_run2.font.name = 'Arial'
    meta_run2.font.size = Pt(10)
    meta_run2.font.color.rgb = NAVY
    
    meta_p.add_run("1.0 (Enterprise Specification)\n")
    
    meta_run3 = meta_p.add_run("Date: ")
    meta_run3.font.bold = True
    meta_run3.font.name = 'Arial'
    meta_run3.font.size = Pt(10)
    meta_run3.font.color.rgb = NAVY
    
    meta_p.add_run("June 2026\n")
    
    meta_run4 = meta_p.add_run("Security: ")
    meta_run4.font.bold = True
    meta_run4.font.name = 'Arial'
    meta_run4.font.size = Pt(10)
    meta_run4.font.color.rgb = NAVY
    
    meta_run4_val = meta_p.add_run("Strictly Private & Confidential")
    meta_run4_val.font.color.rgb = RGBColor(180, 50, 50)
    meta_run4_val.font.bold = True
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Executive Summary & Vision")
    run.font.name = 'Georgia'
    run.font.color.rgb = NAVY
    
    p = doc.add_paragraph(
        "Regent Advisory is a premier luxury real estate brokerage catering to high-net-worth (HNW) "
        "and ultra-high-net-worth (UHNW) individuals, family offices, and institutional investors. In the "
        "ultra-prime property market, clients demand absolute discretion, deep hyper-local market intelligence, "
        "and instantaneous engagement. "
    )
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "This Master Blueprint outlines the architecture, prompts, and technical workflows required to "
        "integrate an advanced, generative AI-powered Real Estate Assistant. Far beyond a generic chatbot, this "
        "assistant acts as an elite digital concierge. It is capable of answering complex property queries, matching "
        "private client profiles to off-market portfolios, validating client requirements, and scheduling high-end, "
        "discrete viewings—all while automatically persisting leads and insights into the Supabase database."
    )
    p.style.font.name = 'Arial'
    
    # ----------------------------------------------------
    # SECTION 2: SYSTEM ARCHITECTURE & INTEGRATION
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    run = h2.add_run("2. System Architecture & Tech Stack")
    run.font.name = 'Georgia'
    run.font.color.rgb = NAVY
    
    p = doc.add_paragraph(
        "To deliver a real-time, highly customized chat experience, the chatbot is architected directly "
        "into the Next.js React ecosystem and integrated with Supabase. The tech stack comprises:"
    )
    
    bullet1 = doc.add_paragraph(style='List Bullet')
    r1 = bullet1.add_run("Core Model Layer: ")
    r1.font.bold = True
    bullet1.add_run("Google Gemini 1.5 Pro / Claude 3.5 Sonnet to ensure unparalleled logical processing, cultural nuance, and professional tone.")
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    r2 = bullet2.add_run("Database Layer (Supabase): ")
    r2.font.bold = True
    bullet2.add_run("Utilizes relational properties and inquiries schemas. Real-time RAG context retrieves matching listings from the public.properties table.")
    
    bullet3 = doc.add_paragraph(style='List Bullet')
    r3 = bullet3.add_run("Vector Engine (pgvector): ")
    r3.font.bold = True
    bullet3.add_run("Allows semantic search of property listings, matches descriptions to client inputs, and surfaces nearby points of interest (schools, private clubs, transits).")
    
    # ----------------------------------------------------
    # SECTION 3: THE PROMPT ENGINE (SYSTEM PROMPTS & TONE)
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    run = h3.add_run("3. Core System Prompt Specifications")
    run.font.name = 'Georgia'
    run.font.color.rgb = NAVY
    
    p = doc.add_paragraph(
        "The AI chatbot's persona is engineered to mimic an elite Mayfair-based private property advisor. "
        "The following structured system prompt enforces strict rules regarding tone, property specifications, "
        "and client discretion:"
    )
    
    # Prompt Box
    prompt_table = doc.add_table(rows=1, cols=1)
    cell = prompt_table.cell(0, 0)
    set_cell_background(cell, "F9FBFD")
    set_cell_margins(cell, top=150, bottom=150, left=200, right=200)
    
    pp = cell.paragraphs[0]
    pp_run = pp.add_run(
        "SYSTEM PROMPT: ELITE PRIVATE REAL ESTATE ADVISOR PERSONA\n\n"
        "ROLE: You are 'Regent Concierge', an ultra-exclusive property advisor representing Regent Advisory in Mayfair, London.\n\n"
        "TONE & STYLE:\n"
        "- Sophisticated, professional, and refined. Use clear, articulate, and grammatically flawless English.\n"
        "- Never use emojis, casual internet abbreviations, or over-the-top sales pitch vocabulary.\n"
        "- Express deep knowledge of high-end areas (e.g., Belgravia, Mayfair, Knightsbridge, Holland Park).\n\n"
        "BEHAVIORAL RULES:\n"
        "1. Discretion First: Under no circumstances expose client names or exact off-market prices unless verified.\n"
        "2. Property Details: Speak accurately regarding square footage, bedroom/bathroom ratios, and premium features.\n"
        "3. Lead Conversion: Gracefully prompt the user to register their details or make an enquiry when they express strong interest or request private viewings.\n"
        "4. Strict Boundary: If a user asks about anything unrelated to real estate, acquisitions, market conditions, or neighborhood features, politely pivot back to luxury real estate advisory services."
    )
    pp_run.font.name = 'Consolas'
    pp_run.font.size = Pt(9.5)
    pp_run.font.color.rgb = CHARCOAL
    
    doc.add_paragraph() # Spacing
    
    # ----------------------------------------------------
    # SECTION 4: CONVERSATIONAL FLOW & SCENARIOS
    # ----------------------------------------------------
    h4 = doc.add_heading(level=1)
    run = h4.add_run("4. Scenario Mapping & Lead Capturing")
    run.font.name = 'Georgia'
    run.font.color.rgb = NAVY
    
    p = doc.add_paragraph(
        "The chatbot maps conversations dynamically to optimize HNW client experience and maximize high-intent "
        "conversions. When certain triggers are hit, specific actions occur:"
    )
    
    # Scenario Table
    table = doc.add_table(rows=4, cols=3)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Client Intent'
    hdr_cells[1].text = 'AI Chatbot Response Strategy'
    hdr_cells[2].text = 'Database Action'
    
    for c in hdr_cells:
        set_cell_background(c, "0C1C30")
        set_cell_margins(c, top=80, bottom=80, left=100, right=100)
        for p_cell in c.paragraphs:
            for run_cell in p_cell.runs:
                run_cell.font.bold = True
                run_cell.font.color.rgb = RGBColor(255, 255, 255)
                run_cell.font.size = Pt(9.5)
                run_cell.font.name = 'Arial'
                
    row_data = [
        ("Acquisitions / Finding properties", "Leverage context vectors to search the 'properties' table. Recommend 2-3 matched properties with square footage and locations.", "Queries 'public.properties'"),
        ("Requesting brochure or floor plans", "Ask for name and email address to verify identity and route the private brochure directly to their inbox.", "Logs to 'public.inquiries'"),
        ("Scheduling private viewings", "Validate dates/times, prompt for phone number/email, and inform them that their dedicated Mayfair advisor will reach out.", "Logs to 'public.inquiries' with property_id linkage")
    ]
    
    for idx, (intent, strategy, action) in enumerate(row_data):
        row_cells = table.rows[idx+1].cells
        row_cells[0].text = intent
        row_cells[1].text = strategy
        row_cells[2].text = action
        
        bg_color = "F9FBFD" if idx % 2 == 0 else "FFFFFF"
        for c in row_cells:
            set_cell_background(c, bg_color)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            for p_cell in c.paragraphs:
                for run_cell in p_cell.runs:
                    run_cell.font.size = Pt(9)
                    run_cell.font.name = 'Arial'
                    run_cell.font.color.rgb = CHARCOAL

    doc.add_paragraph() # Spacing
    
    # ----------------------------------------------------
    # SECTION 5: IMPLEMENTATION & ROLLOUT ROADMAP
    # ----------------------------------------------------
    h5 = doc.add_heading(level=1)
    run = h5.add_run("5. Phased Implementation Roadmap")
    run.font.name = 'Georgia'
    run.font.color.rgb = NAVY
    
    roadmap_steps = [
        ("Phase 1: Database & RAG Configuration (Weeks 1-2)", "Establish pgvector extensions in Supabase. Store full realistic properties records including location metrics, and write SQL function handlers to return matched vector distances."),
        ("Phase 2: Prompt Engineering & UI Assembly (Weeks 3-4)", "Integrate custom client-side chat interfaces styled with premium dark modes and gold accents. Rigorously refine system prompts to avoid any hallucinations or tone deviations."),
        ("Phase 3: Security Audits & DISCRETION Protocols (Week 5)", "Test defenses against prompt injection. Ensure client telephone, emails, and financial information are sanitized or encrypted, and only service-role authenticated handlers store/retrieve critical fields."),
        ("Phase 4: Launch & Human Advisory Handoff (Week 6)", "Deploy live onto the corporate website. Enable seamless desktop and mobile viewings, linking AI logs directly into the company's real-time CRM notification channels.")
    ]
    
    for step_title, step_desc in roadmap_steps:
        p_step = doc.add_paragraph()
        r_step = p_step.add_run(f"■ {step_title}\n")
        r_step.font.bold = True
        r_step.font.color.rgb = GOLD
        r_step.font.name = 'Arial'
        r_step.font.size = Pt(11)
        
        r_desc = p_step.add_run(step_desc)
        r_desc.font.name = 'Arial'
        r_desc.font.size = Pt(10.5)
        r_desc.font.color.rgb = CHARCOAL
        
    doc.add_paragraph("\n" * 2)
    
    # Footer Statement
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run("--- End of Blueprint Document — Proprietary Property of Regent Advisory ---")
    f_run.font.size = Pt(9)
    f_run.font.italic = True
    f_run.font.color.rgb = MUTED_GRAY
    
    # Save document
    output_path = r"e:\Real-State-Advisory-Firm-Demo\Real_Estate_Prompt_Driven_AI_Chatbot_Master_Blueprint.docx"
    doc.save(output_path)
    print(f"Master Blueprint document successfully generated at: {output_path}")

if __name__ == "__main__":
    generate_blueprint()
