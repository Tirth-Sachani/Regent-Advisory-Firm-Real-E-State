import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
    return '\n'.join(fullText)

if __name__ == "__main__":
    content = read_docx(r"e:\Real-State-Advisory-Firm-Demo\Regent_Estate_AI_Master_Implementation_Specification.docx")
    with open(r"e:\Real-State-Advisory-Firm-Demo\scratch\regent_docx_content.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("Done!")
